from docx import Document
from app.table import table_column_widths
from docx.shared import Pt, Inches
from docx.enum.text import WD_BREAK


def excel_to_word(df, new_df, word_file):

    # Create a Word document
    doc = Document()

    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Series")
    run.font.size = Pt(12)
    run.bold = True

    # Add the DataFrame 'df' as a table to the Word document
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style='Table Grid')
    table_column_widths(table, (Inches(2), Inches(5.5),))
    
    # Populate the table with 'df' data
    for i, column in enumerate(df.columns):
        table.cell(0, i).text = column  # Set column headers
        for j in range(df.shape[0]):
            value = df.iloc[j, i]
            cell = table.cell(j + 1, i)
            cell.text = str(value)
            # Set font size for each cell in the table
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(6)  # Set font size to 6 points
            # Reduce spacing within cell
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Set font size and spacing for table header
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(6)  # Set font size to 6 points
                run.font.bold = True
        # Add spacing to table header
        cell.paragraphs[0].paragraph_format.space_before = Pt(6)
        cell.paragraphs[0].paragraph_format.space_after = Pt(6)


    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    # Set table width to match entire page width
    table.autofit = False
    section = doc.sections[-1]
    table_width = section.page_width - section.left_margin - section.right_margin
    table.width = table_width

    paragraph = doc.add_paragraph()
    run = paragraph.add_run("SKUs")
    run.font.size = Pt(12)
    run.bold = True

    # Counter for tables added
    tables_added = 0
    for i in range(1, new_df.shape[1]):
        # Add a paragraph break between tables
        #doc.add_paragraph()

        # Add a new table for the current column
        table = doc.add_table(rows=new_df.shape[0] + 1, cols=2,  style='Table Grid')
        table_column_widths(table, (Inches(1.5), Inches(6),))

        # Populate the table with the values from the first and current column
        for j in range(new_df.shape[0]):
            # Set column headers
            table.cell(0, 0).text = new_df.columns[0]
            table.cell(0, 1).text = new_df.columns[i]

            # Set values for current row
            table.cell(j + 1, 0).text = str(new_df.iloc[j, 0])  # First column value
            table.cell(j + 1, 1).text = str(new_df.iloc[j, i])  # Current column value

            # Set font size for each cell in the table
            for cell in table.rows[j + 1].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(6)  # Set font size to 6 points
                        
                # Reduce spacing within cell
                cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        # Set font size and spacing for table header
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)  # Set font size to 6 points
                    run.font.bold = True
            # Add spacing to table header
            cell.paragraphs[0].paragraph_format.space_before = Pt(6)
            cell.paragraphs[0].paragraph_format.space_after = Pt(6)

        # Set table width to match entire page width
        table.autofit = False
        table.width = table_width

        # Increment tables_added counter
        tables_added += 1

        # Insert a line break after each table
        doc.add_paragraph()

        # Insert a page break after every third table
        if tables_added % 4 == 0:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # Set margins
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.5)  # Set left margin to 0.5 inches
        section.right_margin = Inches(0.5)  # Set right margin to 0.5 inches
        section.top_margin = Inches(0.5)  # Set top margin to 0.5 inches
        section.bottom_margin = Inches(0.5)  # Set bottom margin to 0.5 inches

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.LINE)
    run = paragraph.add_run("Disclaimer: \n This document contains confidential information and is intended only for the product development content review. Do not disseminate, distribute, or copy this document or any of its contents. You cannot use or forward the file without the Product Management team consent or Hewlett-Packard Development Company, L.P. permission. The information contained herein is subject to change without notice. HP shall not be liable for technical or editorial errors or omissions contained herein.")               
    run.font.size = Pt(8)
    run.bold = True

    # Save the Word document
    doc.save(word_file)

